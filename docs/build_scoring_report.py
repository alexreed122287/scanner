#!/usr/bin/env python3
"""Generate the OPTION PANDA scoring system report in PDF + DOCX.

Run:
    python3 docs/build_scoring_report.py

Outputs:
    docs/OPTION_PANDA_Scoring_System.pdf
    docs/OPTION_PANDA_Scoring_System.docx

Source of truth: scoreIt() and STRATEGY_MODES in index.html. When those
change, update the structured data dicts below and rebuild. This script
is hand-authored — there is no autogen from index.html — because the
narrative + efficacy commentary belongs alongside the raw rule weights.
"""
from datetime import date
from pathlib import Path

OUT_DIR = Path(__file__).resolve().parent
PDF_PATH = OUT_DIR / "OPTION_PANDA_Scoring_System.pdf"
DOCX_PATH = OUT_DIR / "OPTION_PANDA_Scoring_System.docx"

# ─── Content (single source of truth for both formats) ──────────────────────

TITLE = "OPTION PANDA — Scoring System & Strategy Guide"
SUBTITLE = "A complete walkthrough of how the dashboard rates every ticker, " \
           "what each rule measures, and how the strategy presets combine them."
VERSION_LINE = f"v2.18.18 build · {date.today().isoformat()}"

EXEC_SUMMARY = [
    "OPTION PANDA scores every ticker in its ~2,500-symbol universe against a "
    "library of 25+ technical, fundamental, options-flow, and sector rules. "
    "Each rule contributes a fixed point weight when its condition fires. "
    "The total is clamped to a 0–164 band; tickers at or above the user-tuned "
    "Min Score threshold are marked GO.",
    "Five built-in strategy presets layer 'must-pass' filters on top of the "
    "score so the user can shift the scanner from broad discovery (any GO) "
    "to evidence-gated rarity (HIGH CONVICTION). Each preset is also paired "
    "with its own minimum score, calibrated against the observed distribution "
    "(peaks 140–160 on a ~164-pt practical max as of the 2026-04-30 retune).",
    "The scoring weights are NOT arbitrary. Every rule was derived from a "
    "named academic source, an industry-standard methodology, or a council-"
    "reviewed backtest. The most heavily-weighted rules (JT 12-1, "
    "Minervini Trend Template, Pocket Pivot, U/D Volume Ratio) are the ones "
    "with the strongest replicated edge in the literature.",
    "This document covers: the full rule catalog (math + source + weight), "
    "the GO-decision logic and ranking tiebreakers, every strategy preset "
    "with its filter set and intent, calibration history, and known "
    "limitations.",
]

# ─── Full rule catalog ──────────────────────────────────────────────────────
# Each tuple: (Rule name, Section, Weight, What it tests, Math / source)
# Section codes:
#   A = Technical Momentum (Tradier bulk quote derived)
#   B = FMP / Fundamentals (analyst data)
#   C = GEX Confirmation (options flow)
#   D = Momentum Acceleration (Tradier history derived)
#   E = Evidence-Based Swing Rules (Tradier history derived)
#   P = Penalty / Overlay (PROVEN, AVOID, sector, broken trend)

RULES = [
    # Overlays / hard adjustments
    ("PROVEN ticker", "P", +8,
     "Ticker is on the curated PROVEN good-mover list.",
     "Hand-maintained list. +8 flat bonus."),
    ("AVOID ticker", "P", -15,
     "Ticker is on the curated AVOID list (chronic chop / scams / fundamentals).",
     "Hand-maintained list. -15 flat penalty and forced NO-GO."),

    # Section A — Technical Momentum
    ("RSI 40-70", "A", +6,
     "Permissive momentum band — not oversold, not extreme.",
     "RSI(14) ≥ 40 and ≤ 70."),
    ("RSI 50-65 (Cardwell)", "A", +6,
     "Cardwell range-shift zone — early entry inside a bull regime.",
     "RSI(14) ≥ 50 and ≤ 65. Bonus on top of RSI 40-70."),
    ("MACD Hist +", "A", +6,
     "MACD histogram > 0 — momentum is positive.",
     "Standard 12/26/9 MACD. macdHist > 0."),
    ("ADX > 25", "A", +6,
     "Confirmed trend strength (Wilder's threshold).",
     "ADX(14) > 25."),
    ("ADX 15-25 + DI+>DI-", "A", +6,
     "Emerging-trend Wilder original signal — DI+ leading DI- with ADX rising.",
     "ADX(14) ∈ [15, 25] and DI+ > DI-. Fires on detail-load only."),
    ("EMA20 > EMA50", "A", +10,
     "Short-term trend above medium-term trend.",
     "EMA(20) > EMA(50) on daily closes."),
    ("TF Aligned 2/3", "A", +12,
     "At least 2 of 3 timeframes (daily / hourly / 15-min) point the same direction.",
     "Tradier multi-timeframe quote aggregation. Count of aligned tf ≥ 2."),
    ("Sector PF > 2", "A", +5,
     "Ticker's static sector profit-factor > 2 in long-call backtest.",
     "SECTOR_PF lookup table from prior backtests."),
    ("52Wk Hi Prox <15%", "A", +8,
     "Price within 15% of the 52-week high — universe-pass for breakout setups.",
     "price ≥ week52hi × 0.85."),
    ("52Wk Hi Prox <5% (G-H)", "A", +6,
     "George-Hwang JF 2004 effect: returns concentrate within 5% of 52-wk hi.",
     "price ≥ week52hi × 0.95. Bonus on top of 15% rule."),
    ("RS > SPY (21d)", "A", +10,
     "21-day relative strength positive vs SPY. Falls back to 5-day if 21d missing.",
     "Total return(21d, ticker) − Total return(21d, SPY) > 0. "
     "5d fallback only when 21d unavailable."),
    ("Price > EMA200", "A", +6,
     "Above the long-term moving average — long-trend bias.",
     "price > EMA(200). Falls back to EMA20 > EMA50 if EMA200 missing."),

    # Section B — FMP
    ("Analyst Revisions ↑", "B", +12,
     "≥2 analysts have raised estimates and there are recent upgrades.",
     "FMP analyst estimates: revisionsUp ≥ 2 AND upgradesLast30 > 0. "
     "Long-calls bias: upward revisions outperform downward revisions "
     "(asymmetry — weighted higher than original +12 here, see commit log)."),
    ("Analyst PT Exists", "B", +8,
     "A consensus analyst price target is published.",
     "FMP price target: ptConsensus present. Light fundamental confirmation."),

    # Section C — GEX (tiered + extreme-positioning penalty)
    ("GEX Call Flow (55–69%)", "C", +4,
     "Emerging call-side OI bias in the 30-DTE chain (constructive but not crowded).",
     "GEX call% of OI ∈ [55, 69]."),
    ("GEX Call Flow (70–74%)", "C", +6,
     "Strong but not overcrowded call-side flow.",
     "GEX call% of OI ∈ [70, 74]."),
    ("GEX Call Flow (≥75%)", "C", -6,
     "Overcrowded call-side OI — contrarian signal, dealers pin/fade.",
     "GEX call% of OI ≥ 75. PENALTY (council 3/3 finding: extreme positioning "
     "is reliably negative on a 1-week horizon)."),

    # Section D — Momentum Acceleration
    ("MFI(14) > 50", "D", +10,
     "Money Flow Index above midline — rising volume-weighted momentum.",
     "MFI(14) > 50 (RSI mid-line analog with volume). "
     "Council 2026-04-29: 14-period is Chande standard; 7 was too noisy."),
    ("CMO(9) > 50", "D", +7,
     "Chande Momentum Oscillator > 50 (Chande's original 1993 threshold).",
     "CMO(9) > 50. 9-period is Chande-documented swing horizon."),

    # Section E — Evidence-Based Swing Rules
    ("Minervini Trend Template", "E", +15,
     "Stage 2 trend confirmation — 6 of 8 SEPA criteria computable from daily bars.",
     "All six must pass: price > MA50; price > MA150; price > MA200; "
     "MA50 > MA150 > MA200; MA200 rising vs 22 bars ago; ≥30% above 52-wk low. "
     "Source: Mark Minervini 'Trade Like a Stock Market Wizard' (2013) SEPA."),
    ("Pocket Pivot", "E", +15,
     "O'Neil / Morales-Kacher institutional accumulation signal.",
     "Today's volume > largest down-day volume of prior 10 days AND close > MA10 "
     "AND close in upper half of day's range AND ≤5% extended from MA10."),
    ("U/D Volume Ratio (50d)", "E", "+12 / +8 / +4",
     "Up-volume vs down-volume ratio over trailing 50 sessions (O'Neil IBD).",
     "Σ(volume on up-close days) / Σ(volume on down-close days) over 50d. "
     "≥2.0 = +12 (strong sponsorship); ≥1.5 = +8 (accumulation); "
     "≥1.25 = +4 (mild); <1.25 = 0."),
    ("VCP — ATR Contraction", "E", +12,
     "Volatility Contraction Pattern: supply drying up before breakout.",
     "ATR(14) today / ATR(14) 30 sessions ago < 0.70. Minervini methodology."),
    ("BBW Squeeze (<p20 of 120d)", "E", +10,
     "Bollinger Band-width compressed in the bottom 20% of its trailing range.",
     "BB(20, 2σ) width / mean. Today's value ≤ 20th-percentile of last 120 days."),
    ("20d Donchian Breakout + Vol", "E", +12,
     "Turtle Traders breakout — Dennis / Eckhardt 1983-88 methodology.",
     "today close > 20-day high AND today volume ≥ 1.5 × 50-day avg volume."),
    ("JT 12-1 Momentum", "E", +20,
     "Cross-sectional momentum, skip last month (Jegadeesh-Titman 1993).",
     "Return from t-252 to t-21 (12 months, skipping recent month). "
     "Threshold: ≥ +25% (top-quintile proxy). "
     "Single strongest replicated factor in equity literature."),

    # Penalties / Overlays
    ("Strong Sector", "P", "+8 / -8",
     "Daily + intraday theme RS overlay vs SPY.",
     "Combined daily + intraday theme score. ≥75 = +8 (top quartile); "
     "≤25 = -8 (bottom quartile). Intraday rotation can trigger same-day."),
    ("Broken Trend (penalty)", "P", -5,
     "Price below EMA50 — medium-term trend broken; calls fight gravity.",
     "price < EMA(50). Penalty only — no positive credit (already implicit "
     "in EMA20>EMA50 and Price>EMA200 rules)."),
]

SECTION_NAMES = {
    "A": "Section A — Technical Momentum (Tradier bulk quote)",
    "B": "Section B — Fundamentals / FMP",
    "C": "Section C — GEX / Options Flow",
    "D": "Section D — Momentum Acceleration (history-based)",
    "E": "Section E — Evidence-Based Swing Rules (history-based)",
    "P": "Overlays & Penalties",
}

# ─── Strategy presets ───────────────────────────────────────────────────────
# Each: (key, label, min_score, must_pass, intent, expected_hit_rate, efficacy_notes)

STRATEGIES = [
    {
        "key": "highConviction",
        "label": "★ HIGH CONVICTION",
        "min_score": 140,
        "must_pass": [
            "PROVEN ticker",
            "Analyst Revisions ↑",
            "GEX Call Heavy (≥55%)",
            "JT 12-1 Momentum (≥25%, top-quintile proxy)",
            "Minervini Trend Template (6/6)",
        ],
        "intent": (
            "Rarest, highest-quality setups. Stacks the two highest-IC scoring "
            "rules (JT 12-1 at +20 and Minervini at +15) as MUST-PASS instead "
            "of merely contributing to the score. A ticker can score 160+ on "
            "other rules but still fail this preset if either trend confirmation "
            "is missing."
        ),
        "expected": "0–5 matches per scan in normal market conditions.",
        "efficacy": (
            "Designed for swing entries where the user wants maximum signal "
            "confluence before paying premium. Backed by: JT 12-1 = strongest "
            "replicated cross-sectional factor (Asness-Moskowitz-Pedersen 2013); "
            "Minervini SEPA = Stage 2 trend filter (Minervini 2013); analyst "
            "revisions = upward earnings surprise lead indicator (Stickel 1991, "
            "Bradshaw 2004); PROVEN ticker bias = curated list of names with "
            "favorable historical long-call distribution."
        ),
    },
    {
        "key": "breakout",
        "label": "▲ BREAKOUT",
        "min_score": 125,
        "must_pass": [
            "52Wk Hi Prox <15%",
            "RS > SPY",
            "MACD Hist +",
            "TF Aligned 2/3",
            "20d Donchian Breakout + Vol",
            "VCP — ATR Contraction",
        ],
        "intent": (
            "Chart-pattern-driven breakout buying with momentum + leadership. "
            "Donchian and VCP are the two highest-IC pure-pattern breakout rules; "
            "requiring BOTH filters out fake-out names that just look extended "
            "without the underlying volatility-contraction or volume-confirmation "
            "signal."
        ),
        "expected": "5–25 matches per scan, market-regime dependent.",
        "efficacy": (
            "Two-mode confirmation: (i) chart-level breakout via 20-day Donchian "
            "high with 1.5× volume (Turtle Traders methodology, Dennis-Eckhardt "
            "1983-88); (ii) volatility contraction prior to the break (Minervini "
            "VCP — ATR(14) ratio < 0.70). Combined with relative strength + "
            "MACD positive + multi-timeframe alignment to screen out reversion-"
            "to-mean candidates."
        ),
    },
    {
        "key": "momentum",
        "label": "▶ MOMENTUM",
        "min_score": 120,
        "must_pass": [
            "MACD Hist +",
            "Price > EMA200",
            "RS > SPY",
            "Strong Sector",
            "JT 12-1 Momentum",
        ],
        "intent": (
            "Confirmed trending stocks with sector tailwind, gated by the "
            "academic gold-standard 12-1 cross-sectional momentum signal. "
            "Making JT 12-1 MUST-PASS turns this preset from 'trendy' into "
            "'trendy AND riding the strongest published factor'."
        ),
        "expected": "5–30 matches per scan during trending markets, fewer in chop.",
        "efficacy": (
            "Primary edge is the 12-1 cross-sectional momentum (Jegadeesh-Titman "
            "1993, Asness-Moskowitz-Pedersen 2013). Layered with sector rotation "
            "(theme overlay top quartile) and standard trend confirmation "
            "(MACD, EMA200, RS). Lower min-score (120) reflects looser breakout "
            "requirement vs. BREAKOUT preset — more inclusive of mid-trend "
            "continuations."
        ),
    },
    {
        "key": "bestItmCalls",
        "label": "◆ BEST ITM CALLS",
        "min_score": 145,
        "must_pass": [
            "MACD Hist +",
            "Price > EMA200",
            "RS > SPY",
            "JT 12-1 Momentum",
            "Minervini Trend Template",
            "Pocket Pivot",
            "Price ≥ $5",
            "Avg Volume ≥ 1,000,000",
            "Vol ratio ≥ 100% (today vs avg)",
            "Max option price ≤ $8",
        ],
        "intent": (
            "Deep-ITM (δ 0.70–0.90) calls on names that pass ALL THREE of the "
            "highest-IC scoring rules. You're paying real premium for deep-ITM "
            "contracts — only setups confirmed by 12-month momentum + Stage 2 "
            "trend template + institutional accumulation should graduate."
        ),
        "expected": "0–8 matches per scan; very selective.",
        "efficacy": (
            "Highest-signal preset by construction. Designed to surface names "
            "where (a) you can afford the contract (≤$8 ceiling), (b) liquidity "
            "is real (≥$5 underlying, ≥1M avg volume, today's volume ≥ average), "
            "and (c) the trend / momentum / accumulation signals all agree. "
            "Pocket Pivot is the institutional-buying flag — distinguishes "
            "true accumulation from just-rallying-on-air."
        ),
    },
    {
        "key": "preMarket",
        "label": "▲ PRE/POST MARKET",
        "min_score": 140,
        "must_pass": [
            "PROVEN ticker",
            "Analyst Revisions ↑",
            "GEX Call Heavy (≥55%)",
            "JT 12-1 Momentum",
            "Minervini Trend Template",
            "Sort: |gap %| descending (overrides standard score sort)",
        ],
        "intent": (
            "Extended-hours gappers screened with the SAME filters as HIGH "
            "CONVICTION, then re-sorted by absolute gap %. Use any time the "
            "extended-hours session is active or on weekends."
        ),
        "expected": "0–5 matches; same selectivity as HIGH CONVICTION, "
                    "re-ordered by gap size.",
        "efficacy": (
            "Behaviorally important: the user explicitly wanted to filter out "
            "random gappers and ONLY surface big movers that also pass the "
            "high-conviction gate. Tradier's `change_percentage` reflects "
            "extended-hours quotes during 3–8:30am CST (pre) and 3–7pm CST "
            "(post); on weekends it shows the most recent extended-hours print."
        ),
    },
]

# ─── Impact & Efficacy by Rule ──────────────────────────────────────────────
# Each tuple:
#   (Rule name, Evidence tier, Expected hit rate, Expected impact (wt × hit),
#    Source / citation, Efficacy notes)
#
# Evidence tier:
#   STRONG    — peer-reviewed factor with replicated edge across decades.
#   MODERATE  — named industry methodology, smaller body of independent evidence.
#   HEURISTIC — project-specific curation or expert prior; no peer review.
#
# Hit rates are empirical estimates from observing the scanner's GO-eligible
# universe (~2,500 names) during the 2026-04 → 2026-05 retune. Treat them as
# ballparks; precise per-rule hit rates shift with regime.

IMPACT_EFFICACY = [
    # Overlays / hard adjustments
    ("PROVEN ticker", "HEURISTIC", "~25–35%", "+2.0 to +2.8",
     "Hand-curated list of names with historically favorable long-call P&L.",
     "Bias-by-design — surfaces names the user has already vetted. Risk: "
     "overfit to the curator's past trades. Mitigated by capping the bonus at +8."),
    ("AVOID ticker", "HEURISTIC", "~3–6%", "−0.5 to −0.9",
     "Hand-curated AVOID list (chronic chop / scams / failed fundamentals).",
     "Hard exclusion — AVOID-flagged tickers are NEVER GO regardless of score. "
     "Lower false-positive cost vs. letting a known-bad name appear in the list."),

    # Section A — Technical Momentum
    ("RSI 40-70", "MODERATE", "~45–55%", "+2.7 to +3.3",
     "Wilder 1978; permissive momentum band, not oversold not extreme.",
     "Broadly-fires baseline filter. IC alone is low; serves as a regime "
     "gate so subsequent stricter rules see a sane subset."),
    ("RSI 50-65 (Cardwell)", "MODERATE", "~20–28%", "+1.2 to +1.7",
     "Andrew Cardwell 'Trading with Range Shift' (1980s) — bullish-regime band.",
     "Adds a tight-zone bonus on top of RSI 40-70. Cardwell's range-shift "
     "framework has shown a small but persistent edge in technical literature; "
     "best when combined with trend filters (which it is, here)."),
    ("MACD Hist +", "MODERATE", "~40–55%", "+2.4 to +3.3",
     "Appel 1979; MACD histogram positive = upward momentum acceleration.",
     "Standard momentum confirmation. Used widely in IBD / O'Neil methodology. "
     "IC modest standalone, useful as part of a multi-rule composite."),
    ("ADX > 25", "MODERATE", "~25–35%", "+1.5 to +2.1",
     "Wilder 1978 — confirmed trend strength threshold.",
     "The Wilder threshold for 'trending' vs. 'choppy'. Filters out range-bound "
     "names where trend rules don't apply."),
    ("ADX 15-25 + DI+>DI-", "MODERATE", "~10–18%", "+0.6 to +1.1",
     "Wilder 1978 — emerging-trend variant of ADX rule.",
     "Bonus for catching trends BEFORE they're confirmed. Higher false-positive "
     "rate than ADX > 25, so weighted equally (+6) and gated by DI+ leading."),
    ("EMA20 > EMA50", "MODERATE", "~45–55%", "+4.5 to +5.5",
     "Classic moving-average crossover (multiple sources).",
     "Short-term trend confirmation. Decent stand-alone IC because it captures "
     "intermediate trend; +10 weight reflects its relative importance vs. RSI."),
    ("TF Aligned 2/3", "HEURISTIC", "~35–55%", "+4.2 to +6.6",
     "Project-specific: multi-timeframe alignment (daily / hourly / 15-min).",
     "No direct academic citation for this exact combination; based on common "
     "trader practice of multi-timeframe confluence. +12 weight assumes 2/3 "
     "alignment is meaningfully more selective than 1/3."),
    ("Sector PF > 2", "HEURISTIC", "~25–35%", "+1.3 to +1.8",
     "SECTOR_PF lookup table built from prior long-call backtest.",
     "Light bias toward sectors where long calls historically worked. Capped "
     "at +5 because backtest n is small; treat as a soft prior."),
    ("52Wk Hi Prox <15%", "STRONG", "~20–35%", "+1.6 to +2.8",
     "George & Hwang Journal of Finance 2004 — '52-Week High and Momentum Investing'.",
     "Replicated factor across markets and decades. Near-52-wk-hi names show "
     "consistent forward outperformance; effect concentrates within 5% of high "
     "(see G-H bonus rule below)."),
    ("52Wk Hi Prox <5% (G-H)", "STRONG", "~6–12%", "+0.4 to +0.7",
     "George & Hwang 2004 — tight-zone bonus where the effect is strongest.",
     "Strong evidence the closer you are to the 52-wk high, the higher the "
     "forward return — paradoxically, NOT mean-reversion. Bonus weighted +6 "
     "to layer on top of the +8 wider band."),
    ("RS > SPY (21d)", "STRONG", "~45–55%", "+4.5 to +5.5",
     "Jegadeesh-Titman 1993 / Carhart 1997 — relative-strength factor.",
     "21-day chosen over 5-day per Lehmann 1990 short-term reversal evidence "
     "(5-day is anti-signal in liquid names). Persistent factor; weighted +10."),
    ("Price > EMA200", "STRONG", "~50–60%", "+3.0 to +3.6",
     "Faber 2007 'A Quantitative Approach to Tactical Asset Allocation'.",
     "Above-200-DMA filter cut max drawdown ~50% in Faber's backtest while "
     "preserving most of the upside. Long-trend bias; baseline filter for many "
     "professional CTAs."),

    # Section B — FMP
    ("Analyst Revisions ↑", "STRONG", "~10–20%", "+1.2 to +2.4",
     "Stickel 1991; Bradshaw 2004 — analyst revisions predict short-term drift.",
     "Upward analyst revisions reliably precede upward price drift over 30–90 days. "
     "Bumped +12 → effective higher weight because long-calls bias asymmetry "
     "(upside is unbounded; the rule captures bullish-side asymmetry)."),
    ("Analyst PT Exists", "HEURISTIC", "~60–80%", "+4.8 to +6.4",
     "Light fundamental confirmation; consensus PT is published.",
     "Mostly a 'is this name covered' filter, not a strong predictive signal. "
     "Weighted +8 to keep it light. Useful as a sanity check — names without "
     "analyst coverage tend to be illiquid or speculative."),

    # Section C — GEX
    ("GEX Call Flow (55–69%)", "MODERATE", "~15–25%", "+0.6 to +1.0",
     "Dealer-positioning literature (SqueezeMetrics, Spotgamma research).",
     "Emerging call-side OI bias — dealers short gamma, pinning effect on the "
     "way up. Constructive but not crowded."),
    ("GEX Call Flow (70–74%)", "MODERATE", "~5–10%", "+0.3 to +0.6",
     "Dealer-positioning literature — strong-but-not-crowded zone.",
     "Sweet spot for long-call setups. Most aggressive dealer-hedging zone "
     "without yet tipping into mean-reversion territory."),
    ("GEX Call Flow (≥75%)", "MODERATE", "~3–7%", "−0.2 to −0.4",
     "Council finding — extreme positioning reliably mean-reverts on 1-week.",
     "PENALTY tier. Above 75% call-OI is crowded; dealers pin or fade. "
     "Tested in council 2026-04-29 backtest as a reliable contrarian signal."),

    # Section D — Momentum Acceleration
    ("MFI(14) > 50", "MODERATE", "~35–50%", "+3.5 to +5.0",
     "Quong-Soudack 1989 — Money Flow Index; volume-weighted RSI analog.",
     "Volume-confirmed momentum. Standard 14-period; threshold 50 mirrors "
     "RSI mid-line logic. Bumped from 7-period after Connors research showed "
     "shorter periods are mostly noise."),
    ("CMO(9) > 50", "MODERATE", "~18–28%", "+1.3 to +2.0",
     "Chande 1993 — Chande Momentum Oscillator at Chande's documented threshold.",
     "Stricter than MFI — fewer hits but each pass is a stronger signal. "
     "Chande's original 1993 paper documented the >50 threshold; later "
     "literature uses zero-line cross which is too permissive."),

    # Section E — Evidence-Based Swing Rules
    ("Minervini Trend Template", "STRONG", "~5–12%", "+0.8 to +1.8",
     "Minervini 'Trade Like a Stock Market Wizard' (2013) — SEPA Stage 2.",
     "Six computable criteria. Strict (~5–12% pass rate) but each pass "
     "confirms Stage 2 trend. Strong out-of-sample evidence in Minervini's "
     "verified trading records and IBD-style methodology."),
    ("Pocket Pivot", "STRONG", "~3–8%", "+0.5 to +1.2",
     "O'Neil / Morales-Kacher 'How to Trade in Stocks' (2010).",
     "Institutional-accumulation flag — today's volume beats the largest "
     "down-day volume of prior 10. Designed to catch buying BEFORE the "
     "obvious breakout. Documented edge in O'Neil/Morales backtests."),
    ("U/D Volume Ratio (50d)", "STRONG", "tiered ~30–50%", "+1.2 to +6.0",
     "O'Neil 'How to Make Money in Stocks' (1988) — IBD methodology.",
     "Sum of up-day volume vs. down-day volume over 50d. Replicated across "
     "decades by IBD; ≥1.5 indicates accumulation. Tiered weight rewards "
     "stronger accumulation: +4 / +8 / +12 for 1.25 / 1.5 / 2.0+ ratios."),
    ("VCP — ATR Contraction", "STRONG", "~5–10%", "+0.6 to +1.2",
     "Minervini 2013 — Volatility Contraction Pattern (Stage 2 precursor).",
     "ATR(14) today / ATR(14) 30d ago < 0.70 indicates supply contraction "
     "before a breakout. Strict by design — rare but high-conviction."),
    ("BBW Squeeze (<p20 of 120d)", "MODERATE", "~20%", "+2.0",
     "Bollinger 1980s — Bollinger Bandwidth compression precedes expansion.",
     "By construction this fires for the bottom-20% of trailing 120 days, so "
     "hit rate ≈ 20%. The 'precedes expansion' edge is modest; weight +10 "
     "reflects that it's a regime indicator, not a directional one."),
    ("20d Donchian Breakout + Vol", "STRONG", "~3–8%", "+0.4 to +1.0",
     "Dennis-Eckhardt 1983-88 — Turtle Traders methodology.",
     "Close above 20-day high with 1.5×+ volume. Best-documented systematic "
     "trend-following edge in equity literature. Strict — fires rarely but "
     "each pass is a high-confidence breakout."),
    ("JT 12-1 Momentum", "STRONG", "~10–20%", "+2.0 to +4.0",
     "Jegadeesh-Titman 1993; Asness-Moskowitz-Pedersen 2013.",
     "Strongest replicated factor in equity markets. Return from t-252 to "
     "t-21 (12 months, skip last month). +25% threshold approximates the "
     "top quintile of S&P 500 momentum. Highest single-rule weight (+20)."),

    # Penalties / Overlays
    ("Strong Sector", "MODERATE", "~25% each tail", "+2.0 / −2.0",
     "Sector rotation literature; daily + intraday theme RS overlay.",
     "Top-quartile theme = +8, bottom-quartile = −8. Long-calls strategy "
     "depends on sector tailwind more than market-neutral strategies do, "
     "which is why this was bumped ±5 → ±8."),
    ("Broken Trend (penalty)", "MODERATE", "~30–40%", "−1.5 to −2.0",
     "Trend-following literature; medium-term trend break flag.",
     "Price below EMA50 = calls fight gravity. Penalty only (no positive "
     "credit — already implicit in other trend rules). −5 weight is light "
     "to avoid double-counting the trend signal."),
]

# Strategy-level efficacy summary — qualitative assessment of how each preset
# combines its underlying rules and what its expected efficacy profile is.
# Used as the "Strategy-Level Efficacy" subsection after the per-rule table.

STRATEGY_EFFICACY = [
    {
        "name": "★ HIGH CONVICTION",
        "ic_factors": (
            "JT 12-1 (highest replicated IC factor in equities), Minervini "
            "Trend (Stage 2 trend confirmation), Analyst Revisions ↑ (lead "
            "indicator), PROVEN bias, GEX call flow."
        ),
        "selectivity": "Highest — typically 0–5 matches per scan.",
        "expected_win_rate": (
            "Unknown empirically (small n=19 backtest); theoretically the "
            "highest of any preset given the must-pass stack of three of the "
            "strongest factors in the literature. Behavioral risk: small "
            "sample → high variance regardless of edge."
        ),
        "best_for": (
            "Patient-capital swing entries where the user wants maximum "
            "signal confluence and is willing to wait for 0-match days."
        ),
    },
    {
        "name": "▲ BREAKOUT",
        "ic_factors": (
            "20d Donchian (Turtle Traders), VCP — ATR contraction (Minervini), "
            "52-wk high proximity (George-Hwang), RS > SPY, MACD."
        ),
        "selectivity": "Medium — typically 5–25 matches per scan.",
        "expected_win_rate": (
            "Donchian breakouts have the longest replicated track record in "
            "trend-following systems. Pattern-based; weaker IC than factor-"
            "based momentum but cleaner entry timing. Best in trending "
            "markets, fails frequently in chop."
        ),
        "best_for": (
            "Active traders who want chart-level breakout entries with both "
            "volatility-contraction confirmation (VCP) and volume-confirmation "
            "(Donchian 1.5×) instead of bare price-only signals."
        ),
    },
    {
        "name": "▶ MOMENTUM",
        "ic_factors": (
            "JT 12-1 (gold-standard cross-sectional momentum), RS > SPY "
            "(intermediate momentum), Strong Sector (rotation overlay), "
            "MACD + EMA200 (trend confirmation)."
        ),
        "selectivity": "Medium — typically 5–30 matches per scan.",
        "expected_win_rate": (
            "Gold-standard preset by IC. JT 12-1 alone has 30+ years of "
            "out-of-sample evidence; layering with sector rotation + trend "
            "confirmation should improve risk-adjusted returns vs. pure "
            "momentum. Weaker in market regime shifts (early 2008, "
            "Q1 2020, 2022 H1)."
        ),
        "best_for": (
            "Trend-followers who want the academic-edge backbone (12-1 "
            "momentum) gated by simple trend filters. Sweet spot in "
            "persistent uptrends."
        ),
    },
    {
        "name": "◆ BEST ITM CALLS",
        "ic_factors": (
            "Stacks all three highest-IC rules (JT 12-1, Minervini, Pocket "
            "Pivot) + liquidity floors + max option price ceiling."
        ),
        "selectivity": "Very high — typically 0–8 matches per scan.",
        "expected_win_rate": (
            "Theoretically highest per-trade win rate given the triple-"
            "confirmation stack, but smallest sample size. Designed for the "
            "specific tactic of buying deep-ITM (δ 0.70–0.90) calls where "
            "you're paying real premium and need maximum signal confluence."
        ),
        "best_for": (
            "Users who want to deploy a meaningful amount of premium per "
            "name and need every signal aligned before doing so. Liquidity "
            "floors (≥$5 / ≥1M avg vol) prevent slippage in execution."
        ),
    },
    {
        "name": "▲ PRE/POST MARKET",
        "ic_factors": (
            "Same factor stack as HIGH CONVICTION but re-sorted by |gap %|."
        ),
        "selectivity": "Very high — typically 0–5 matches.",
        "expected_win_rate": (
            "Surfaces gaps that ALSO pass the high-conviction stack — "
            "avoiding the well-known 'gap-and-fade' problem where random "
            "earnings gaps mean-revert. Limited evidence for extended-hours "
            "edge in liquid names; treat as a session-specific overlay "
            "rather than a standalone factor."
        ),
        "best_for": (
            "Trading the pre-market or after-hours session when a vetted "
            "name has gapped meaningfully (>2–3%). Avoids the random-"
            "gapper trap."
        ),
    },
]

# ─── GO decision + tiebreakers ──────────────────────────────────────────────

GO_DECISION = [
    "After all rules fire, the raw score is clamped to the 0–164 band. "
    "Penalties (Broken Trend −5, weak sector −8, AVOID −15) can drive the "
    "subtotal negative; the clamp forces it to 0 so the threshold comparator "
    "doesn't see nonsensical scores.",
    "GO threshold is read from the Min Score input on the scanner tab. "
    "Default 140 was calibrated 2026-04-30 against the observed distribution "
    "(peaks 140–160 on a ~164-pt practical max). Lowering the threshold "
    "widens the GO list; raising it tightens. Strategy presets override the "
    "default with their own per-strategy min-score (see preset table).",
    "AVOID-list tickers can NEVER be GO regardless of score — the AVOID flag "
    "forces NO-GO downstream of the score clamp.",
]

TIEBREAKERS = [
    ("1. Score", "Primary sort — descending."),
    ("2. techScore", "Sum of all Section A + D + E rules. Stronger technical "
                    "setup wins when total scores tie."),
    ("3. relStr5d", "5-day relative-strength % vs SPY. Stronger short-term "
                   "momentum wins."),
    ("4. Theme dailyScore", "Sector / theme strength from theme_scores.json. "
                            "Stronger sector context wins."),
    ("5. Volume", "Daily volume — liquidity floor as the final tie-break."),
]

# ─── Calibration & efficacy ─────────────────────────────────────────────────

CALIBRATION = [
    ("2026-04-29 council overhaul (drops)",
     "Removed three anti-signal rules: RSI(7) Accel (≈70% noise per Connors), "
     "5d %chg Accel (Lehmann 1990 short-term reversal), Vol > Avg flat threshold "
     "(too crude — replaced by U/D Volume Ratio)."),
    ("2026-04-29 council overhaul (retunes)",
     "Tightened existing rules to evidence-based parameters: RSI band → Cardwell "
     "range-shift; MFI(7) → MFI(14) > 50; CMO zero-cross → CMO(9) > 50; "
     "5d RS → 21d RS; 52-wk Hi <15% → <15% pass + <5% George-Hwang bonus."),
    ("2026-04-29 council overhaul (additions)",
     "Added 7 evidence-based rules: Minervini Template (+15), Pocket Pivot (+15), "
     "U/D Volume Ratio (+4/+8/+12), VCP (+12), BBW Squeeze (+10), Donchian (+12), "
     "JT 12-1 Momentum (+20). Practical max raised from ~110 to ~164."),
    ("2026-04-30 GO threshold retune",
     "Default GO threshold raised 100 → 150, then re-calibrated 2026-05-03 to "
     "140 against the observed distribution. New rules rarely hit max "
     "simultaneously — distribution peak sits 140–160, not 165+."),
    ("2026-05-01 GEX retune",
     "GEX Call Flow rule moved from flat +6 to tiered with extreme-positioning "
     "PENALTY: 55–69% = +4 (emerging), 70–74% = +6 (strong), ≥75% = −6 "
     "(overcrowded contrarian). Backed by council finding that ≥75% OI on calls "
     "reliably mean-reverts on a 1-week horizon."),
    ("2026-05-04 Analyst Revisions bump",
     "Analyst Revisions ↑ bumped +12 → +16 for the long-calls-only user. "
     "Upward revisions outperform downward revisions asymmetrically — call-only "
     "strategies should weight bullish signals more heavily."),
    ("2026-05-03 Strong Sector bump",
     "Strong Sector ±5 → ±8. Bullish unhedged bets depend more heavily on "
     "sector tailwind than hedged or directional-pair strategies do."),
]

KNOWN_LIMITATIONS = [
    "Single-leg long-call bias. The scoring is explicitly tuned for long-call "
    "swing trades. It is not validated for short premium, spreads, puts, or "
    "longer-than-50-DTE positions.",
    "Backtest n=19 trades (2026-04-29 council review). The empirical edge of "
    "the combined system is small and not statistically separated from noise. "
    "Two-thirds of cumulative P&L came from one outlier (POET 12.5C, 8-DTE, "
    "+$1,225). Treat all rule weights as expert-judgment-with-priors, not "
    "empirically-validated weights.",
    "Some rules require ≥210 daily bars (Minervini, Donchian, U/D Volume, BBW). "
    "Fresh IPOs or newly-uplisted symbols are flagged with dataGap = true and "
    "treated as 'unknown', NOT as failed. Strategies that require those rules "
    "MUST-PASS will simply exclude such names.",
    "GEX data is sourced from a separate options-flow tab (runs once per scan + "
    "scheduled at 9:35 / 12:00 / 15:00 ET on weekdays). If GEX hasn't run yet "
    "the GEX Call Flow rule scores 0 — neither bonus nor penalty. Strategies "
    "that MUST-PASS GEX will exclude all tickers until GEX data lands.",
    "Cross-sectional ranking (JT 12-1's true form) is approximated via a hard "
    "+25% threshold rather than ranking against the universe simultaneously. "
    "This is a simplification — true cross-sectional momentum would require "
    "computing returns for every scanned ticker before classifying. Threshold "
    "is set near the historical top-quintile cut.",
    "Theme / sector overlay depends on theme_scores.json which is regenerated "
    "by a nightly GitHub Actions cron. If the cron fails (network blip, "
    "Tradier outage, etc.), the Strong Sector rule pass-throughs (no bonus / "
    "no penalty) so the rest of the scoring still works.",
]

# ─── End-user playbook ──────────────────────────────────────────────────────

USER_PLAYBOOK = [
    ("1. Set your API keys",
     "Open the API tab and enter your Tradier proxy URL + live-mode token, "
     "and your FMP API key. Without an FMP key, Section B rules score 0 — "
     "the rest still works."),
    ("2. Pick a strategy preset",
     "Tap STRATEGY at the top of the scanner. Start with HIGH CONVICTION "
     "to see the most selective gate, or BREAKOUT / MOMENTUM if you want "
     "more candidates. CLEAR returns to the default Min Score 140 with no "
     "must-pass filters."),
    ("3. Run the scan",
     "Tap ↻ SCAN. The bulk pass scores all ~2,500 tickers in 30–60s. The "
     "foreground enrichment pass adds Tradier history + FMP data to the "
     "top 60. A background pass extends Tradier history to ranks 61–200 "
     "over the next ~5 minutes; you'll see scores update live."),
    ("4. Read the results table",
     "Sortable columns: Score, Tech, Fund, Price, Change %, Theme, Sector, "
     "GEX flow, etc. Click any row to open the detail pane with the full "
     "rule-by-rule breakdown — every passing rule with its point value, "
     "every failing rule with what it would have scored."),
    ("5. Use Export Matrix or Export GOs",
     "EXPORT GOs copies a comma-separated list of every ticker that's GO "
     "under the current filter. EXPORT MATRIX in the Portfolio tab pulls "
     "all positions into a Google Sheet for sharing."),
    ("6. Tune Min Score if needed",
     "If a preset returns 0 matches in a quiet market, drop Min Score by "
     "10–20 to widen the GO band. If it returns 50+ in a hot market, raise "
     "it to tighten. The score distribution shifts with regime."),
]


# ─── PDF builder (reportlab) ────────────────────────────────────────────────

def build_pdf():
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle,
        KeepTogether,
    )
    from xml.sax.saxutils import escape as _xml_escape

    def E(s):
        """Escape <, >, & for reportlab's Paragraph parser."""
        return _xml_escape(str(s), {})

    doc = SimpleDocTemplate(
        str(PDF_PATH), pagesize=LETTER,
        leftMargin=0.7 * inch, rightMargin=0.7 * inch,
        topMargin=0.7 * inch, bottomMargin=0.7 * inch,
        title=TITLE, author="Option Panda",
    )
    styles = getSampleStyleSheet()
    title_st = ParagraphStyle(
        "title", parent=styles["Title"], fontSize=22, spaceAfter=8,
        textColor=colors.HexColor("#0b4a7a"))
    subtitle_st = ParagraphStyle(
        "subtitle", parent=styles["Heading3"], fontSize=11, textColor=colors.grey,
        leading=14, spaceAfter=4, alignment=TA_CENTER)
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], fontSize=16,
                        textColor=colors.HexColor("#0b4a7a"), spaceBefore=12,
                        spaceAfter=6)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], fontSize=13,
                        textColor=colors.HexColor("#0b4a7a"), spaceBefore=10,
                        spaceAfter=4)
    h3 = ParagraphStyle("h3", parent=styles["Heading3"], fontSize=11,
                        textColor=colors.HexColor("#0b4a7a"), spaceBefore=8,
                        spaceAfter=2)
    body = ParagraphStyle("body", parent=styles["BodyText"], fontSize=10,
                          leading=14, spaceAfter=6, alignment=TA_LEFT)
    bullet = ParagraphStyle("bullet", parent=body, leftIndent=14, bulletIndent=0)
    small = ParagraphStyle("small", parent=body, fontSize=9, leading=12,
                           textColor=colors.dimgrey)
    cell_body = ParagraphStyle("cell", parent=body, fontSize=8.5, leading=11,
                               spaceAfter=0)
    cell_bold = ParagraphStyle("cellb", parent=cell_body, fontName="Helvetica-Bold")

    story = []

    # Title page
    story.append(Spacer(1, 1.4 * inch))
    story.append(Paragraph(TITLE, title_st))
    story.append(Spacer(1, 6))
    story.append(Paragraph(SUBTITLE, subtitle_st))
    story.append(Spacer(1, 24))
    story.append(Paragraph(VERSION_LINE, subtitle_st))
    story.append(PageBreak())

    # Executive summary
    story.append(Paragraph("Executive Summary", h1))
    for p in EXEC_SUMMARY:
        story.append(Paragraph(E(p), body))

    # Scoring engine overview
    story.append(Paragraph("1. The Scoring Engine", h1))
    story.append(Paragraph(
        "Every ticker enters scoreIt() and runs through ~25 rules grouped into "
        "five sections plus overlays. Each rule has a fixed point weight, "
        "fires independently, and never short-circuits — a single failing rule "
        "doesn't disqualify a ticker, only reduces its total score.", body))
    story.append(Paragraph(
        "Practical max score = 164 points. Empirical peak in normal markets "
        "= 140–160. Default GO threshold = 140 (~85% of practical max).", body))

    # Rule catalog table
    story.append(Paragraph("2. Full Rule Catalog", h1))
    story.append(Paragraph(
        "Every rule the dashboard ships. Weight, what it tests, and the math / "
        "academic source. Rules grouped by section.", body))

    for section_code, section_name in SECTION_NAMES.items():
        section_rules = [r for r in RULES if r[1] == section_code]
        if not section_rules:
            continue
        story.append(Paragraph(section_name, h2))
        data = [[Paragraph("<b>Rule</b>", cell_bold),
                 Paragraph("<b>Wt</b>", cell_bold),
                 Paragraph("<b>Tests</b>", cell_bold),
                 Paragraph("<b>Math / Source</b>", cell_bold)]]
        for name, _sec, weight, tests, math in section_rules:
            data.append([
                Paragraph(E(name), cell_body),
                Paragraph(E(weight), cell_body),
                Paragraph(E(tests), cell_body),
                Paragraph(E(math), cell_body),
            ])
        tbl = Table(data, colWidths=[1.5*inch, 0.5*inch, 2.0*inch, 2.8*inch],
                    repeatRows=1)
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dbe7f3")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0b4a7a")),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8.5),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd6e2")),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1),
             [colors.white, colors.HexColor("#f5f8fb")]),
        ]))
        story.append(tbl)
        story.append(Spacer(1, 6))

    # Impact & efficacy by rule
    story.append(PageBreak())
    story.append(Paragraph("3. Impact &amp; Efficacy by Rule", h1))
    story.append(Paragraph(
        "Per-rule view of how often each scoring element fires, how much it "
        "contributes to the total score on average, and the evidence tier "
        "behind it. Evidence tiers: <b>STRONG</b> = peer-reviewed factor "
        "with replicated edge across decades; <b>MODERATE</b> = named "
        "industry methodology, smaller body of independent evidence; "
        "<b>HEURISTIC</b> = project-specific curation or expert prior, no "
        "peer review.", body))
    story.append(Paragraph(
        "Hit-rate ballparks are empirical estimates from the 2026-04 → 2026-05 "
        "retune across the ~2,500-symbol universe. Precise per-rule rates "
        "shift with market regime. Expected impact = rule weight × hit rate.",
        body))

    ie_data = [[
        Paragraph("<b>Rule</b>", cell_bold),
        Paragraph("<b>Evidence</b>", cell_bold),
        Paragraph("<b>Hit %</b>", cell_bold),
        Paragraph("<b>Avg Impact</b>", cell_bold),
        Paragraph("<b>Source / Efficacy</b>", cell_bold),
    ]]
    for name, tier, hit, impact, source, notes in IMPACT_EFFICACY:
        ie_data.append([
            Paragraph(E(name), cell_body),
            Paragraph(E(tier), cell_body),
            Paragraph(E(hit), cell_body),
            Paragraph(E(impact), cell_body),
            Paragraph(E(source) + "<br/>" + E(notes), cell_body),
        ])
    ie_tbl = Table(ie_data,
                   colWidths=[1.3*inch, 0.7*inch, 0.7*inch, 0.7*inch, 3.4*inch],
                   repeatRows=1)
    ie_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dbe7f3")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd6e2")),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1),
         [colors.white, colors.HexColor("#f5f8fb")]),
    ]))
    story.append(ie_tbl)
    story.append(Spacer(1, 8))

    story.append(Paragraph("Strategy-level efficacy", h2))
    story.append(Paragraph(
        "Each preset is a curated combination of the rules above. The "
        "efficacy of a preset depends on which underlying rules dominate "
        "its must-pass stack — presets gated by <b>STRONG</b>-tier rules "
        "inherit those rules' evidence base.", body))
    for s in STRATEGY_EFFICACY:
        story.append(Paragraph(E(s["name"]), h3))
        story.append(Paragraph("<b>IC-driving factors.</b> " + E(s["ic_factors"]), body))
        story.append(Paragraph("<b>Selectivity.</b> " + E(s["selectivity"]), body))
        story.append(Paragraph("<b>Expected win rate.</b> " + E(s["expected_win_rate"]), body))
        story.append(Paragraph("<b>Best for.</b> " + E(s["best_for"]), body))

    # GO decision + tiebreakers
    story.append(PageBreak())
    story.append(Paragraph("4. GO Decision &amp; Ranking", h1))
    story.append(Paragraph("How the GO flag is set", h2))
    for p in GO_DECISION:
        story.append(Paragraph(E(p), body))

    story.append(Paragraph("Tiebreaker order", h2))
    tb_data = [[Paragraph("<b>Priority</b>", cell_bold),
                Paragraph("<b>Field</b>", cell_bold)]]
    for name, desc in TIEBREAKERS:
        tb_data.append([Paragraph(E(name), cell_body), Paragraph(E(desc), cell_body)])
    tb_tbl = Table(tb_data, colWidths=[1.4*inch, 5.4*inch], repeatRows=1)
    tb_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dbe7f3")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd6e2")),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1),
         [colors.white, colors.HexColor("#f5f8fb")]),
    ]))
    story.append(tb_tbl)
    story.append(Spacer(1, 6))

    # Strategy presets
    story.append(PageBreak())
    story.append(Paragraph("5. Strategy Presets", h1))
    story.append(Paragraph(
        "Five preset combinations layer must-pass filters on top of the score, "
        "plus their own minimum score threshold. Click any preset chip in the "
        "STRATEGY row to apply.", body))

    for s in STRATEGIES:
        story.append(KeepTogether([
            Paragraph(E(f"{s['label']} — Min Score {s['min_score']}"), h2),
            Paragraph("<b>Intent.</b> " + E(s["intent"]), body),
            Paragraph("<b>Must-pass filters.</b>", body),
        ]))
        for f in s["must_pass"]:
            story.append(Paragraph("&bull; " + E(f), bullet))
        story.append(Paragraph("<b>Expected matches.</b> " + E(s["expected"]), body))
        story.append(Paragraph("<b>Efficacy / rationale.</b> " + E(s["efficacy"]), body))
        story.append(Spacer(1, 6))

    # Strategy comparison table
    story.append(Paragraph("Strategy comparison at a glance", h2))
    comp_data = [[
        Paragraph("<b>Preset</b>", cell_bold),
        Paragraph("<b>Min Score</b>", cell_bold),
        Paragraph("<b># Must-Pass</b>", cell_bold),
        Paragraph("<b>Expected hits</b>", cell_bold),
    ]]
    for s in STRATEGIES:
        comp_data.append([
            Paragraph(E(s["label"]), cell_body),
            Paragraph(E(s["min_score"]), cell_body),
            Paragraph(E(len([f for f in s["must_pass"]
                              if not f.lower().startswith("sort:")])), cell_body),
            Paragraph(E(s["expected"]), cell_body),
        ])
    comp_tbl = Table(comp_data, colWidths=[1.8*inch, 0.8*inch, 1.0*inch, 3.2*inch],
                     repeatRows=1)
    comp_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dbe7f3")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd6e2")),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1),
         [colors.white, colors.HexColor("#f5f8fb")]),
    ]))
    story.append(comp_tbl)

    # Calibration history
    story.append(PageBreak())
    story.append(Paragraph("6. Calibration History", h1))
    story.append(Paragraph(
        "The rule set and thresholds were derived through a structured review "
        "process (\"council\" — domain advisors for quant, microstructure, "
        "risk, behavioral, engineering). Each entry below is a dated change "
        "to the scoring system with the rationale.", body))
    for date_label, desc in CALIBRATION:
        story.append(Paragraph("<b>" + E(date_label) + "</b>", h3))
        story.append(Paragraph(E(desc), body))

    # Efficacy + limitations
    story.append(Paragraph("7. Efficacy & Known Limitations", h1))
    story.append(Paragraph(
        "Honest framing: the rule weights and thresholds are expert-judgment-"
        "informed-by-literature, not empirically backed by a large-sample "
        "backtest of THIS COMBINED SYSTEM. Treat the score as a structured "
        "checklist, not a guaranteed edge.", body))
    for p in KNOWN_LIMITATIONS:
        story.append(Paragraph("&bull; " + E(p), bullet))

    # User playbook
    story.append(PageBreak())
    story.append(Paragraph("8. End-User Playbook", h1))
    story.append(Paragraph(
        "Step-by-step for someone opening the dashboard for the first time.",
        body))
    for step, desc in USER_PLAYBOOK:
        story.append(Paragraph(E(step), h3))
        story.append(Paragraph(E(desc), body))

    # Appendix
    story.append(PageBreak())
    story.append(Paragraph("Appendix — Glossary", h1))
    glossary = [
        ("GO / NO-GO", "Binary classification output by scoreIt. GO = score ≥ "
                       "Min Score AND not AVOID-listed."),
        ("Min Score", "User-tunable input on the scanner tab (default 140). "
                      "Strategy presets override with their own per-strategy "
                      "minimum."),
        ("Must-Pass", "Filter applied AFTER scoring. A ticker that scores 200 "
                      "but fails a must-pass filter (e.g. no Analyst Revisions ↑) "
                      "is excluded from the preset's GO list."),
        ("techScore", "Sum of all Section A + D + E rule weights that fired for "
                      "this ticker. Used as the second tiebreaker."),
        ("fundScore", "Sum of Section B + Strong Sector. Surfaces analyst + "
                      "sector tailwind contribution."),
        ("GEX", "Gamma exposure — dealer-hedging pressure derived from the 30-DTE "
                "options chain. Call Heavy = ≥55% of OI on the call side."),
        ("JT 12-1", "Jegadeesh-Titman 12-1 momentum — return from t-252 to t-21 "
                    "(skip last month)."),
        ("VCP", "Volatility Contraction Pattern — Minervini's ATR-shrinkage "
                "precursor to breakout."),
        ("Donchian", "Turtle Traders breakout — close above prior 20-day high "
                     "on 1.5×+ volume."),
        ("BBW", "Bollinger Bandwidth — (upper − lower) / mean. Compressed when "
                "in the bottom 20% of the trailing 120-day range."),
        ("PROVEN ticker", "Hand-curated list of names with historically favorable "
                          "long-call P&L. +8 bonus on score."),
        ("AVOID ticker", "Hand-curated AVOID list. -15 penalty + forced NO-GO."),
        ("Theme overlay", "Sector/theme RS from theme_scores.json (nightly cron). "
                          "Drives the Strong Sector ±8 rule."),
        ("Intraday RS", "Live theme RS computed from member changePct vs SPY "
                        "during the day so rotation triggers same-day."),
    ]
    g_data = [[Paragraph("<b>Term</b>", cell_bold),
               Paragraph("<b>Definition</b>", cell_bold)]]
    for term, defn in glossary:
        g_data.append([Paragraph(E(term), cell_body), Paragraph(E(defn), cell_body)])
    g_tbl = Table(g_data, colWidths=[1.4*inch, 5.4*inch], repeatRows=1)
    g_tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dbe7f3")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd6e2")),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1),
         [colors.white, colors.HexColor("#f5f8fb")]),
    ]))
    story.append(g_tbl)

    story.append(Spacer(1, 18))
    story.append(Paragraph(
        "Source of truth: scoreIt() and STRATEGY_MODES in index.html. "
        "When those change, update build_scoring_report.py and rebuild.", small))

    doc.build(story)


# ─── DOCX builder (python-docx) ─────────────────────────────────────────────

def build_docx():
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    # Base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    BRAND = RGBColor(0x0B, 0x4A, 0x7A)

    def add_title(text, size=22, color=BRAND, center=True):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(size)
        run.font.color.rgb = color
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_h(text, level=1):
        sizes = {1: 15, 2: 12, 3: 11}
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(sizes.get(level, 11))
        run.font.color.rgb = BRAND

    def add_body(text, bold_prefix=None):
        p = doc.add_paragraph()
        if bold_prefix:
            run = p.add_run(bold_prefix)
            run.bold = True
        p.add_run(text)

    def add_bullet(text):
        doc.add_paragraph(text, style="List Bullet")

    def shade(cell, color_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), color_hex)
        tcPr.append(shd)

    def add_table(headers, rows, col_widths=None):
        t = doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Light Grid Accent 1"
        hdr_cells = t.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = ""
            p = hdr_cells[i].paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(9.5)
            shade(hdr_cells[i], "DBE7F3")
        for r_idx, r in enumerate(rows):
            cells = t.rows[r_idx + 1].cells
            for c_idx, c in enumerate(r):
                cells[c_idx].text = ""
                p = cells[c_idx].paragraphs[0]
                run = p.add_run(str(c))
                run.font.size = Pt(9)
                if r_idx % 2 == 1:
                    shade(cells[c_idx], "F5F8FB")
        if col_widths:
            for i, w in enumerate(col_widths):
                for row in t.rows:
                    row.cells[i].width = w

    # Title page
    doc.add_paragraph()
    doc.add_paragraph()
    add_title(TITLE, size=22)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(SUBTITLE)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(VERSION_LINE)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    doc.add_page_break()

    # Executive summary
    add_h("Executive Summary", level=1)
    for p_txt in EXEC_SUMMARY:
        add_body(p_txt)

    # Engine overview
    add_h("1. The Scoring Engine", level=1)
    add_body("Every ticker enters scoreIt() and runs through ~25 rules grouped "
             "into five sections plus overlays. Each rule has a fixed point "
             "weight, fires independently, and never short-circuits — a single "
             "failing rule doesn't disqualify a ticker, only reduces its total "
             "score.")
    add_body("Practical max score = 164 points. Empirical peak in normal "
             "markets = 140–160. Default GO threshold = 140 (~85% of practical "
             "max).")

    # Rule catalog
    add_h("2. Full Rule Catalog", level=1)
    add_body("Every rule the dashboard ships. Weight, what it tests, and the "
             "math / academic source. Grouped by section.")
    for sc, sn in SECTION_NAMES.items():
        srules = [r for r in RULES if r[1] == sc]
        if not srules:
            continue
        add_h(sn, level=2)
        rows = [(r[0], str(r[2]), r[3], r[4]) for r in srules]
        add_table(["Rule", "Wt", "Tests", "Math / Source"], rows,
                  col_widths=[Inches(1.5), Inches(0.5), Inches(2.0), Inches(2.8)])

    # Impact & efficacy by rule
    doc.add_page_break()
    add_h("3. Impact & Efficacy by Rule", level=1)
    add_body("Per-rule view of how often each scoring element fires, how much "
             "it contributes to the total score on average, and the evidence "
             "tier behind it. Evidence tiers: STRONG = peer-reviewed factor "
             "with replicated edge across decades; MODERATE = named industry "
             "methodology, smaller body of independent evidence; HEURISTIC "
             "= project-specific curation or expert prior, no peer review.")
    add_body("Hit-rate ballparks are empirical estimates from the 2026-04 → "
             "2026-05 retune across the ~2,500-symbol universe. Precise "
             "per-rule rates shift with market regime. Expected impact = "
             "rule weight × hit rate.")
    ie_rows = []
    for name, tier, hit, impact, source, notes in IMPACT_EFFICACY:
        ie_rows.append((name, tier, hit, impact, source + "\n" + notes))
    add_table(["Rule", "Evidence", "Hit %", "Avg Impact", "Source / Efficacy"],
              ie_rows,
              col_widths=[Inches(1.3), Inches(0.7), Inches(0.7), Inches(0.7),
                          Inches(3.4)])

    add_h("Strategy-level efficacy", level=2)
    add_body("Each preset is a curated combination of the rules above. The "
             "efficacy of a preset depends on which underlying rules dominate "
             "its must-pass stack — presets gated by STRONG-tier rules "
             "inherit those rules' evidence base.")
    for s in STRATEGY_EFFICACY:
        add_h(s["name"], level=3)
        add_body(s["ic_factors"], bold_prefix="IC-driving factors. ")
        add_body(s["selectivity"], bold_prefix="Selectivity. ")
        add_body(s["expected_win_rate"], bold_prefix="Expected win rate. ")
        add_body(s["best_for"], bold_prefix="Best for. ")

    # GO decision
    doc.add_page_break()
    add_h("4. GO Decision & Ranking", level=1)
    add_h("How the GO flag is set", level=2)
    for p_txt in GO_DECISION:
        add_body(p_txt)
    add_h("Tiebreaker order", level=2)
    add_table(["Priority", "Field"], TIEBREAKERS,
              col_widths=[Inches(1.4), Inches(5.4)])

    # Strategy presets
    doc.add_page_break()
    add_h("5. Strategy Presets", level=1)
    add_body("Five preset combinations layer must-pass filters on top of the "
             "score, plus their own minimum score threshold. Click any preset "
             "chip in the STRATEGY row to apply.")
    for s in STRATEGIES:
        add_h(f"{s['label']} — Min Score {s['min_score']}", level=2)
        add_body(s["intent"], bold_prefix="Intent. ")
        add_body("", bold_prefix="Must-pass filters.")
        for f in s["must_pass"]:
            add_bullet(f)
        add_body(s["expected"], bold_prefix="Expected matches. ")
        add_body(s["efficacy"], bold_prefix="Efficacy / rationale. ")

    add_h("Strategy comparison at a glance", level=2)
    comp_rows = []
    for s in STRATEGIES:
        comp_rows.append((
            s["label"], str(s["min_score"]),
            str(len([f for f in s["must_pass"]
                    if not f.lower().startswith("sort:")])),
            s["expected"],
        ))
    add_table(["Preset", "Min Score", "# Must-Pass", "Expected hits"], comp_rows,
              col_widths=[Inches(1.8), Inches(0.8), Inches(1.0), Inches(3.2)])

    # Calibration
    doc.add_page_break()
    add_h("6. Calibration History", level=1)
    add_body("The rule set and thresholds were derived through a structured "
             "review process (\"council\" — domain advisors for quant, "
             "microstructure, risk, behavioral, engineering). Each entry below "
             "is a dated change to the scoring system with the rationale.")
    for date_label, desc in CALIBRATION:
        add_h(date_label, level=3)
        add_body(desc)

    # Limitations
    add_h("7. Efficacy & Known Limitations", level=1)
    add_body("Honest framing: the rule weights and thresholds are expert-"
             "judgment-informed-by-literature, not empirically backed by a "
             "large-sample backtest of THIS COMBINED SYSTEM. Treat the score as "
             "a structured checklist, not a guaranteed edge.")
    for p_txt in KNOWN_LIMITATIONS:
        add_bullet(p_txt)

    # Playbook
    doc.add_page_break()
    add_h("8. End-User Playbook", level=1)
    add_body("Step-by-step for someone opening the dashboard for the first "
             "time.")
    for step, desc in USER_PLAYBOOK:
        add_h(step, level=3)
        add_body(desc)

    # Appendix
    doc.add_page_break()
    add_h("Appendix — Glossary", level=1)
    glossary = [
        ("GO / NO-GO", "Binary classification output by scoreIt. GO = score ≥ "
                       "Min Score AND not AVOID-listed."),
        ("Min Score", "User-tunable input on the scanner tab (default 140). "
                      "Strategy presets override with their own per-strategy "
                      "minimum."),
        ("Must-Pass", "Filter applied AFTER scoring. A ticker that scores 200 "
                      "but fails a must-pass filter is excluded from the "
                      "preset's GO list."),
        ("techScore", "Sum of all Section A + D + E rule weights that fired. "
                      "Second tiebreaker."),
        ("fundScore", "Sum of Section B + Strong Sector. Surfaces analyst + "
                      "sector tailwind contribution."),
        ("GEX", "Gamma exposure — dealer-hedging pressure derived from the "
                "30-DTE options chain. Call Heavy = ≥55% of OI on calls."),
        ("JT 12-1", "Jegadeesh-Titman 12-1 momentum: return from t-252 to t-21."),
        ("VCP", "Volatility Contraction Pattern — Minervini's ATR-shrinkage "
                "precursor to breakout."),
        ("Donchian", "Turtle Traders breakout — close above prior 20-day high "
                     "on 1.5×+ volume."),
        ("BBW", "Bollinger Bandwidth — compressed when in the bottom 20% of "
                "the trailing 120-day range."),
        ("PROVEN ticker", "Hand-curated list with historically favorable long-"
                          "call P&L. +8 bonus."),
        ("AVOID ticker", "Hand-curated AVOID list. -15 + forced NO-GO."),
        ("Theme overlay", "Sector / theme RS from theme_scores.json (nightly "
                          "cron). Drives Strong Sector ±8."),
        ("Intraday RS", "Live theme RS computed from member changePct vs SPY "
                        "during the day so rotation triggers same-day."),
    ]
    add_table(["Term", "Definition"], glossary,
              col_widths=[Inches(1.4), Inches(5.4)])

    p = doc.add_paragraph()
    run = p.add_run("Source of truth: scoreIt() and STRATEGY_MODES in "
                    "index.html. When those change, update "
                    "build_scoring_report.py and rebuild.")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(str(DOCX_PATH))


if __name__ == "__main__":
    build_pdf()
    build_docx()
    print(f"Wrote {PDF_PATH}")
    print(f"Wrote {DOCX_PATH}")
